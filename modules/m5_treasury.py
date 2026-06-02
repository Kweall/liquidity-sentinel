import requests
import pandas as pd
import numpy as np
import os
import re
import glob
from bs4 import BeautifulSoup
from urllib3.exceptions import InsecureRequestWarning
from docx import Document
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))
from modules.utils import mad_score, validate_module_output

requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

RAW_DIR = "data/raw"
PROCESSED_DIR = "data/processed"
BASE_URL = "https://www.cbr.ru"
SORS_PAGE = "https://www.cbr.ru/statistics/bank_sector/sors/"

def _get_budget_excel_url() -> str | None:
    try:
        resp = requests.get(SORS_PAGE, headers={'User-Agent': 'Mozilla/5.0'}, timeout=30)
        resp.raise_for_status()
    except Exception:
        return None
    
    soup = BeautifulSoup(resp.text, 'html.parser')
    toggle = soup.find('a', string=lambda t: t and "Бюджетные средства на счетах кредитных организаций" in t)
    if not toggle:
        return None

    controls_id = toggle.get('aria-controls')
    content_div = soup.find('div', id=controls_id) if controls_id else toggle.find_next_sibling('div')
    if not content_div:
        return None

    file_link = content_div.find('a', href=lambda h: h and h.lower().endswith('.xlsx'))
    if not file_link:
        return None

    href = file_link.get('href')
    return BASE_URL + href if href.startswith('/') else href

def fetch_cbr_budget_data() -> pd.DataFrame:
    excel_url = _get_budget_excel_url()
    if not excel_url:
        return pd.DataFrame()
    
    os.makedirs(RAW_DIR, exist_ok=True)
    local_path = os.path.join(RAW_DIR, "cbr_budget_funds.xlsx")

    try:
        resp = requests.get(excel_url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=30)
        resp.raise_for_status()
        with open(local_path, 'wb') as f:
            f.write(resp.content)
    except Exception:
        return pd.DataFrame()

    try:
        df_raw = pd.read_excel(local_path, header=None)
        target_row_idx = None
        for idx, row in df_raw.iterrows():
            if "Остатки бюджетных средств на счетах, всего" in str(row[0]):
                target_row_idx = idx
                break
        if target_row_idx is None:
            for idx, row in df_raw.iterrows():
                if "Остатки бюджетных средств" in str(row[0]):
                    target_row_idx = idx
                    break

        if target_row_idx is None:
            return pd.DataFrame()

        dates_raw = df_raw.iloc[1, 1:].values
        values_raw = df_raw.iloc[target_row_idx, 1:].values

        result = pd.DataFrame({'date': dates_raw, 'total_eks_mln': values_raw})
        result['date'] = pd.to_datetime(result['date'], errors='coerce', dayfirst=True)
        result['total_eks_mln'] = pd.to_numeric(
            result['total_eks_mln'].astype(str).str.replace(' ', '').str.replace(',', '.'), errors='coerce'
        )
        result = result.dropna(subset=['date', 'total_eks_mln'])
        result['total_eks'] = result['total_eks_mln'] / 1000.0
        return result[['date', 'total_eks']].sort_values('date').reset_index(drop=True)
    except Exception:
        return pd.DataFrame()

def fetch_roskazna_deposits(docx_path: str = None) -> pd.DataFrame:
    if docx_path is None:
        files = glob.glob(os.path.join(RAW_DIR, "*.docx"))
        if not files:
            return pd.DataFrame()
        docx_path = files[0]
    
    if not os.path.exists(docx_path):
        return pd.DataFrame()

    try:
        doc = Document(docx_path)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
    except Exception:
        return pd.DataFrame()

    if not lines:
        return pd.DataFrame()

    records = []
    current = {}
    keywords = {
        "дата проведения отбора заявок": "date",
        "общий объем средств в подлежащих удовлетворению заявках": "placed_volume_mln",
        "общий объем направленных заявок": "total_bids_mln",
        "процентная ставка отсечения": "cut_off_rate",
        "количество кредитных организаций, принявших участие": "participants",
        "срок размещения, в днях": "term_days"
    }

    i = 0
    while i < len(lines):
        line = lines[i]
        line_lower = line.lower().strip()

        if "дата проведения отбора заявок" in line_lower:
            if current and "date" in current:
                records.append(current)
            current = {}
            if i + 1 < len(lines):
                val = re.sub(r'[^\d\.]', '', lines[i+1].replace(',', '.')).strip()
                if val:
                    current['date'] = val
                i += 1
        else:
            for key, col in keywords.items():
                if key in line_lower:
                    val = line.split(key)[-1].strip().lstrip(':').strip()
                    if not val and i + 1 < len(lines):
                        val = lines[i+1].strip()
                        i += 1
                    val = re.sub(r'[^\d\.\-]', '', val.replace(',', '.')).strip()
                    if val:
                        current[col] = val
                    break
        i += 1

    if current and "date" in current:
        records.append(current)

    if not records:
        return pd.DataFrame()

    df = pd.DataFrame(records)
    df['date'] = pd.to_datetime(df['date'], dayfirst=True, errors='coerce')
    
    for col in ['placed_volume_mln', 'total_bids_mln', 'cut_off_rate', 'participants', 'term_days']:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')

    if 'placed_volume_mln' in df.columns:
        df['deposits_placed'] = df['placed_volume_mln'] / 1000.0
    if 'total_bids_mln' in df.columns:
        df['total_bids'] = df['total_bids_mln'] / 1000.0

    agg_dict = {'deposits_placed': 'sum'}
    for c in ['total_bids', 'cut_off_rate', 'participants']:
        if c in df.columns:
            agg_dict[c] = 'last'

    df = df.dropna(subset=['date']).groupby('date').agg(agg_dict).reset_index()
    return df[['date'] + [c for c in agg_dict.keys() if c in df.columns]].sort_values('date')

def process_m5(cbr_df: pd.DataFrame, roskazna_df: pd.DataFrame = None) -> pd.DataFrame:
    if cbr_df.empty:
        return pd.DataFrame()

    df = cbr_df.copy()
    
    if roskazna_df is not None and not roskazna_df.empty:
        roskazna_unique = roskazna_df.drop_duplicates(subset=['date'], keep='last')
        df = df.merge(roskazna_unique, on='date', how='left')
        df['deposits_placed'] = df['deposits_placed'].ffill().bfill().fillna(0)
        df['participants'] = df['participants'].ffill().bfill().fillna(0)
    else:
        df['deposits_placed'] = 0
        df['participants'] = 0
    
    df = df.sort_values('date')
    df['delta_eks_monthly'] = df['total_eks'].diff()
    df['delta_deposits'] = df['deposits_placed'].diff() if 'deposits_placed' in df.columns else np.nan
    
    # Порог 10 млрд для оттока (было 20)
    df['Flag_Budget_Drain'] = ((df['delta_eks_monthly'].abs() > 10) & (df['delta_eks_monthly'] < 0)).astype(int)
    
    # Дополнительный флаг для сильного оттока (>30 млрд)
    df['Flag_Strong_Drain'] = ((df['delta_eks_monthly'].abs() > 30) & (df['delta_eks_monthly'] < 0)).astype(int)
    
    df['Flag_EndOfMonth'] = (df['date'] == df['date'].dt.to_period('M').dt.end_time.dt.normalize()).astype(int)

    quarter_end_months = [3, 6, 9, 12]
    df['Flag_EndOfQuarter'] = (df['date'].dt.month.isin(quarter_end_months) & 
                                (df['date'] == df['date'].dt.to_period('Q').dt.end_time.dt.normalize())).astype(int)
    
    if 'deposits_placed' in df.columns:
        df['Flag_HighPlacement'] = (df['deposits_placed'] > df['deposits_placed'].median() * 1.5).astype(int)
    else:
        df['Flag_HighPlacement'] = 0
    
    df['mad_score_cbr'] = mad_score(df['delta_eks_monthly'], window=60)
    df['mad_score_roskazna'] = mad_score(df['delta_deposits'], window=60) if 'delta_deposits' in df.columns else np.nan
    
    cbr_stress = df['mad_score_cbr'].abs().fillna(0)
    roskazna_stress = df['mad_score_roskazna'].abs().fillna(0) if 'mad_score_roskazna' in df.columns else 0
    
    cbr_stress = (cbr_stress * 1.0).clip(0, 10)
    roskazna_stress = (roskazna_stress * 1.0).clip(0, 10)
    
    df['stress_m5'] = (0.7 * cbr_stress + 0.3 * roskazna_stress).clip(0, 10)
    
    df['stress_m5'] *= np.where(df['Flag_Budget_Drain'] == 1, 1.5, 1.0)
    df['stress_m5'] *= np.where(df['Flag_Strong_Drain'] == 1, 2.0, 1.0)
    df['stress_m5'] = df['stress_m5'].clip(0, 10)
    
    result = df[[
        'date',
        'total_eks',
        'delta_eks_monthly',
        'deposits_placed',
        'participants',
        'mad_score_cbr',
        'mad_score_roskazna',
        'stress_m5',
        'Flag_Budget_Drain',
        'Flag_Strong_Drain',
        'Flag_EndOfMonth',
        'Flag_EndOfQuarter',
        'Flag_HighPlacement'
    ]].copy()
    
    result = result.dropna(subset=['mad_score_cbr'])
    result['date'] = pd.to_datetime(result['date']).astype('datetime64[ns]')
    result = result.drop_duplicates(subset=['date']).sort_values('date')
    
    validate_module_output(result, "M5")
    return result

def run(docx_file: str = None) -> pd.DataFrame:
    os.makedirs(RAW_DIR, exist_ok=True)
    os.makedirs(PROCESSED_DIR, exist_ok=True)
    cbr_data = fetch_cbr_budget_data()
    roskazna_data = fetch_roskazna_deposits(docx_file)
    result = process_m5(cbr_data, roskazna_data)
    
    if not result.empty:
        result.to_parquet(os.path.join(PROCESSED_DIR, "m5_output.parquet"), index=False)
        print(f"\nМ5 готов: {len(result)} строк")
        print(f"Период: {result['date'].min().date()} — {result['date'].max().date()}")
        print(f"Stress диапазон: {result['stress_m5'].min():.2f} — {result['stress_m5'].max():.2f}")
        print(f"Флагов оттока (порог 10 млрд): {result['Flag_Budget_Drain'].sum()}")
        print(f"Флагов сильного оттока (30 млрд): {result['Flag_Strong_Drain'].sum()}")
        print(f"Средний stress_m5: {result['stress_m5'].mean():.2f}")
        
    return result

if __name__ == "__main__":
    run()